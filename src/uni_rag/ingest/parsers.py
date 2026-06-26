"""Document parsers for PDF, DOCX, Markdown."""
from __future__ import annotations
import logging
from dataclasses import dataclass
from pathlib import Path
import re
import fitz  # PyMuPDF
import docx
from markdown_it import MarkdownIt

from uni_rag.ingest.mineru_client import is_mineru_available, parse_file_via_api

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Output of a parser."""
    text: str
    format: str  # 'pdf' | 'docx' | 'md'
    source_path: str
    pages: list[tuple[int, str]] | None = None  # [(page_no, text)] for pdf


def _parse_pdf(path: Path) -> ParsedDocument:
    """Try MinerU cloud API first, fall back to PyMuPDF on failure."""
    if is_mineru_available():
        try:
            md_text = parse_file_via_api(path)
            return ParsedDocument(
                text=md_text,
                format="pdf",
                source_path=str(path),
                pages=None,  # MinerU returns flat Markdown; page numbers lost
            )
        except Exception as e:
            logger.warning("MinerU 解析失败，回退到 PyMuPDF: %s", e)
    return _parse_pdf_pymupdf(path)


def _parse_pdf_pymupdf(path: Path) -> ParsedDocument:
    doc = fitz.open(str(path))
    pages = []
    full = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        pages.append((i + 1, text))
        full.append(text)
    doc.close()
    return ParsedDocument(
        text="\n\n".join(full),
        format="pdf",
        source_path=str(path),
        pages=pages,
    )


def _parse_docx(path: Path) -> ParsedDocument:
    d = docx.Document(str(path))

    parts: list[str] = []

    # 1. 段落（按 body 顺序）
    for p in d.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        text = _wrap_formulas(text)
        parts.append(text)

    # 2. 表格（按 body 顺序；用 d.element.body 顺序，简化：放在段落后）
    for table in d.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return ParsedDocument(
        text="\n\n".join(parts),
        format="docx",
        source_path=str(path),
    )


_FORMULA_BLOCK_RE = re.compile(r"\$\$([^$]+)\$\$", re.DOTALL)


def _wrap_formulas(text: str) -> str:
    """把 $$...$$ 块公式包到 ``` 围栏里；$...$ 行内公式保留。"""
    def block_sub(m: re.Match) -> str:
        return f"```\n{m.group(1).strip()}\n```"
    return _FORMULA_BLOCK_RE.sub(block_sub, text)


def _parse_md(path: Path) -> ParsedDocument:
    md = MarkdownIt()
    text = path.read_text(encoding="utf-8")
    rendered = md.render(text)
    # MarkdownIt 渲染后有 HTML 标签，去掉
    import re
    plain = re.sub(r"<[^>]+>", "", rendered)
    plain = re.sub(r"\n{3,}", "\n\n", plain).strip()
    return ParsedDocument(
        text=plain,
        format="md",
        source_path=str(path),
    )


def parse_document(path: str | Path) -> ParsedDocument:
    """Dispatch to the right parser by file extension."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Not found: {p}")
    ext = p.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(p)
    if ext == ".docx":
        return _parse_docx(p)
    if ext in (".md", ".markdown"):
        return _parse_md(p)
    raise ValueError(f"Unsupported format: {ext}")
