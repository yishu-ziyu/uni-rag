"""Generate a DOCX fixture with a 2x2 table and a math-formula paragraph."""
import docx
from pathlib import Path

OUT = Path(__file__).resolve().parent.parent / "tests" / "fixtures" / "sample_with_table_formula.docx"


def main():
    d = docx.Document()

    d.add_heading("Chapter 1: Formulas and Tables", level=1)
    d.add_paragraph("This chapter introduces the quadratic equation and a comparison table.")

    d.add_heading("Section 1.1: Quadratic Formula", level=2)
    d.add_paragraph("The well-known identity is: $x^2 + y^2 = z^2$ (Pythagoras).")
    d.add_paragraph("For integrals we write: $$\\int_0^1 f(x) dx$$ in display form.")

    d.add_heading("Section 1.2: Comparison", level=2)
    d.add_paragraph("The table below compares supervised and unsupervised learning:")

    # 2x2 table
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Aspect"
    table.cell(0, 1).text = "Description"
    table.cell(1, 0).text = "Supervised"
    table.cell(1, 1).text = "Uses labeled data"

    d.add_paragraph("End of document.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
